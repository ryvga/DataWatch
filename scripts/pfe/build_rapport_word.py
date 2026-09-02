from __future__ import annotations

import json
import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/mounir/Documents/Claude/Projects/DataWatch")
SOURCE = ROOT / "docs/pfe/report_source.json"
OUT = ROOT / "output/pfe/Rapport_PFE_DataWatch_Mounir_Gaiby.docx"

FONT = "Arial"
INK = "202124"
MUTED = "5F6368"
LIGHT = "EEF0F3"
ACCENT = "B1202D"
ACCENT_SOFT = "F8EAEC"

FIGURES = {
    "Figure 1.1": ROOT / "docs/diagrams/pfe/gantt-doc.png",
    "Figure 3.1": ROOT / "docs/diagrams/pfe/use_cases-doc.png",
    "Figure 3.2": ROOT / "docs/diagrams/pfe/sequence-uml-doc.png",
    "Figure 3.3": ROOT / "docs/diagrams/pfe/classes-doc.png",
    "Figure 3.4": ROOT / "docs/diagrams/pfe/architecture-doc.png",
    "Figure 4.1": ROOT / "docs/screenshots/pfe/01-operations-report.png",
    "Figure 4.2": ROOT / "docs/screenshots/pfe/02-incident-orders-report.png",
    "Figure 4.3": ROOT / "docs/screenshots/pfe/03-table-orders-report.png",
    "Figure 4.4": ROOT / "docs/screenshots/pfe/04-alerts-report.png",
    "Figure 4.5": ROOT / "docs/screenshots/pfe/05-ai-governance-report.png",
    "Figure 4.6": ROOT / "docs/diagrams/pfe/validation-doc.png",
    "Figure 4.7": ROOT / "docs/diagrams/pfe/demo_flow-doc.png",
}

SKIP_BODY_INDEXES = set(range(185, 190)) | set(range(199, 204)) | set(range(241, 246)) | {224, 228, 232, 236, 238}


def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_border(paragraph, color=ACCENT, size="14", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def remove_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def add_field(paragraph, instruction, fallback=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r1 = OxmlElement("w:r"); r1.append(begin)
    r2 = OxmlElement("w:r"); r2.append(instr)
    r3 = OxmlElement("w:r"); r3.append(separate)
    paragraph._p.extend([r1, r2, r3])
    if fallback:
        run = paragraph.add_run(fallback)
        set_font(run, 10, color=MUTED)
    r4 = OxmlElement("w:r"); r4.append(end)
    paragraph._p.append(r4)


def set_page_numbering(section, fmt, start=1):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))


def clear_container(container):
    for p in list(container.paragraphs):
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)


def set_footer(section, numbered=True):
    section.footer.is_linked_to_previous = False
    clear_container(section.footer)
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    if numbered:
        run = p.add_run("—  ")
        set_font(run, 8.5, color=MUTED)
        add_field(p, "PAGE")
        run = p.add_run("  —")
        set_font(run, 8.5, color=MUTED)


def set_header(section, text=""):
    section.header.is_linked_to_previous = False
    clear_container(section.header)
    p = section.header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        set_font(run, 8, color=MUTED)


def section_geometry(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(1.0)


def update_fields_on_open(doc):
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    h1.font.name = FONT; h1.font.size = Pt(18); h1.font.bold = True; h1.font.color.rgb = RGBColor.from_string(INK)
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT); h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1.paragraph_format.space_before = Pt(6); h1.paragraph_format.space_after = Pt(14)
    h1.paragraph_format.keep_with_next = True; h1.paragraph_format.page_break_before = True

    h2 = styles["Heading 2"]
    h2.font.name = FONT; h2.font.size = Pt(13.5); h2.font.bold = True; h2.font.color.rgb = RGBColor.from_string(INK)
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT); h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.paragraph_format.space_before = Pt(14); h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = FONT; h3.font.size = Pt(11.5); h3.font.bold = True; h3.font.color.rgb = RGBColor.from_string(INK)
    h3._element.rPr.rFonts.set(qn("w:ascii"), FONT); h3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h3.paragraph_format.space_before = Pt(10); h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    cap = styles["Caption"]
    cap.font.name = FONT; cap.font.size = Pt(9); cap.font.italic = False; cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap._element.rPr.rFonts.set(qn("w:ascii"), FONT); cap._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4); cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.keep_with_next = False; cap.paragraph_format.keep_together = True


def add_numbering(doc):
    root = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def create(abstract_id, num_id, fmt, text):
        abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
        lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0"); abstract.append(lvl)
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
        nf = OxmlElement("w:numFmt"); nf.set(qn("w:val"), fmt); lvl.append(nf)
        lt = OxmlElement("w:lvlText"); lt.set(qn("w:val"), text); lvl.append(lt)
        jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab); ppr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360"); ppr.append(ind); lvl.append(ppr)
        rpr = OxmlElement("w:rPr"); fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), FONT); fonts.set(qn("w:hAnsi"), FONT); rpr.append(fonts); lvl.append(rpr)
        root.insert(0, abstract)
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id)); aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id)); num.append(aid); root.append(num)

    create(next_abs, next_num, "bullet", "•")
    create(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid]); ppr.append(num_pr)


def add_body(doc, text, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False, bold=False, after=6):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, 11, bold=bold, italic=italic, color=INK)
    return p


def add_list(doc, text, num_id):
    p = doc.add_paragraph(style="Normal")
    set_num(p, num_id)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, 10.7, color=INK)
    return p


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bookmark_id)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1, start); paragraph._p.append(end)


def add_internal_link(paragraph, text, anchor, bold=False):
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("w:anchor"), anchor); hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle"); rstyle.set(qn("w:val"), "Hyperlink"); rpr.append(rstyle)
    fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), FONT); fonts.set(qn("w:hAnsi"), FONT); rpr.append(fonts)
    size = OxmlElement("w:sz"); size.set(qn("w:val"), "17"); rpr.append(size)
    color = OxmlElement("w:color"); color.set(qn("w:val"), INK); rpr.append(color)
    if bold: rpr.append(OxmlElement("w:b"))
    run.append(rpr); node = OxmlElement("w:t"); node.text = text; run.append(node); hyperlink.append(run); paragraph._p.append(hyperlink)


def add_heading(doc, text, level, front=False, anchor=None, bookmark_id=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if front else WD_ALIGN_PARAGRAPH.LEFT
        if front:
            p.paragraph_format.space_after = Pt(20)
            remove_paragraph_border(p)
        else:
            paragraph_border(p)
    run = p.add_run(text)
    set_font(run, 18 if level == 1 else (13.5 if level == 2 else 11.5), bold=True, color=INK)
    if anchor is not None:
        add_bookmark(p, anchor, bookmark_id)
    return p


def add_live_index(doc, entries, instruction):
    """Add a Word-updatable index with visible, clickable cached results."""
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    doc_part = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery"); gallery.set(qn("w:val"), "Table of Contents")
    unique = OxmlElement("w:docPartUnique")
    doc_part.extend([gallery, unique]); sdt_pr.append(doc_part); sdt.append(sdt_pr)
    content = OxmlElement("w:sdtContent")
    made = []
    for text, page, anchor, level in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.55 if level == 2 else 0)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1.5); p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_internal_link(p, text, anchor, bold=(level == 1))
        r = p.add_run("\t" + str(page)); set_font(r, 8.5, bold=(level == 1), color=INK)
        made.append(p)
    first = made[0]._p
    begin = OxmlElement("w:r"); char = OxmlElement("w:fldChar"); char.set(qn("w:fldCharType"), "begin"); begin.append(char)
    instr_r = OxmlElement("w:r"); instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction; instr_r.append(instr)
    sep_r = OxmlElement("w:r"); sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); sep_r.append(sep)
    first.insert(1, begin); first.insert(2, instr_r); first.insert(3, sep_r)
    end_r = OxmlElement("w:r"); end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); end_r.append(end); made[-1]._p.append(end_r)
    for p in made:
        p._p.getparent().remove(p._p); content.append(p._p)
    sdt.append(content)
    body = doc._body._body
    body.insert(body.index(body.sectPr), sdt)


def add_picture(doc, path, label, caption, anchor=None, bookmark_id=None):
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        w, h = im.size
    max_w = 15.8
    max_h = 15.2
    width = max_w
    height = width * h / w
    if height > max_h:
        height = max_h
        width = height * w / h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Cm(width), height=Cm(height))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    r = cap.add_run(label)
    set_font(r, 9, bold=True, color=ACCENT)
    r = cap.add_run(" — " + caption)
    set_font(r, 9, color=MUTED)
    if anchor is not None:
        add_bookmark(cap, anchor, bookmark_id)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ppr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F3F4F6"); ppr.append(shd)
        run = p.add_run(line)
        set_font(run, 9, color=INK)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")


def set_table_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + tag))
        if node is None:
            node = OxmlElement("w:" + tag); tc_mar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa)); tc_w.set(qn("w:type"), "dxa")


def add_abbreviations(doc):
    rows = [
        ("API", "Application Programming Interface"), ("IA", "Intelligence artificielle"),
        ("LLM", "Large Language Model"), ("SaaS", "Software as a Service"),
        ("SLA", "Service Level Agreement"), ("JWT", "JSON Web Token"),
        ("ORM", "Object Relational Mapping"), ("CI", "Continuous Integration"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")); tbl_w.set(qn("w:type"), "dxa"); tbl_w.set(qn("w:w"), "8950")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in (2200, 6750):
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for i, text in enumerate(("Abréviation", "Signification")):
        c = table.rows[0].cells[i]; set_cell_width(c, (2200, 6750)[i]); set_table_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), ACCENT_SOFT); c._tc.get_or_add_tcPr().append(shd)
        c.text = ""; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); set_font(r, 10, bold=True, color=INK)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader"); header.set(qn("w:val"), "true"); tr_pr.append(header)
    for abbr, meaning in rows:
        cells = table.add_row().cells
        for i, text in enumerate((abbr, meaning)):
            set_cell_width(cells[i], (2200, 6750)[i]); set_table_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = ""; p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text); set_font(r, 10, bold=(i == 0), color=INK)
    # Quiet horizontal rules only.
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders")
            for edge in ("top", "bottom"):
                e = OxmlElement("w:" + edge); e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "D7DADF"); borders.append(e)
            for edge in ("start", "end"):
                e = OxmlElement("w:" + edge); e.set(qn("w:val"), "nil"); borders.append(e)
            tc_pr.append(borders)


def remap_heading(text):
    replacements = {
        "1. Introduction générale": "INTRODUCTION GÉNÉRALE",
        "1.1 Résultats attendus": "Résultats attendus",
        "1.2 Organisation du rapport": "Organisation du rapport",
        "6. Conclusion générale et perspectives": "CONCLUSION GÉNÉRALE ET PERSPECTIVES",
        "6.1 Cadre du travail": "Cadre du travail",
        "6.2 Synthèse des apports": "Synthèse des apports",
        "6.3 Perspectives": "Perspectives",
        "7. Références": "RÉFÉRENCES",
        "8. Annexes": "ANNEXES",
    }
    if text in replacements:
        return replacements[text]
    m = re.match(r"([2-5])\. Chapitre ([1-4]) : (.+)", text)
    if m:
        return f"CHAPITRE {m.group(2)} — {m.group(3).upper()}"
    m = re.match(r"([2-5])\.(\d(?:\.\d)?) (.+)", text)
    if m:
        old = int(m.group(1)); return f"{old-1}.{m.group(2)} {m.group(3)}"
    return text


def cover(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
    logo = p.add_run().add_picture(str(ROOT / "docs/pfe/assets/isga-logo.png"), width=Cm(5.8))
    logo._inline.docPr.set("descr", "Logo officiel de l’ISGA")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ISGA CASABLANCA"); set_font(r, 11, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(16)
    r = p.add_run("3CI — Big Data et Intelligence Artificielle"); set_font(r, 10, color=MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(22)
    r = p.add_run("PROJET DE FIN D’ÉTUDES"); set_font(r, 13, bold=True, color=ACCENT)
    paragraph_border(p, color=ACCENT, size="10", space="10")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Conception et réalisation de DataWatch"); set_font(r, 24, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.right_indent = Cm(1.0); p.paragraph_format.space_after = Pt(34)
    r = p.add_run("Plateforme SaaS multi-tenant de surveillance de la qualité des données, enrichie par l’IA pour l’explication des incidents"); set_font(r, 12.5, color=MUTED)
    entries = [("RÉALISÉ PAR", "Mounir Gaiby"), ("ENCADRÉ PAR", "Dr. HANINE MOHAMED"), ("FILIÈRE", "3CI — Big Data et Intelligence Artificielle")]
    for label, value in entries:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.left_indent = Cm(0.25); p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(6.2), WD_TAB_ALIGNMENT.LEFT)
        r = p.add_run(label); set_font(r, 8.5, bold=True, color=ACCENT)
        r = p.add_run("\t" + value); set_font(r, 10.5, bold=True, color=INK)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Année universitaire 2025–2026"); set_font(r, 10, color=MUTED)


def is_bullet_index(index):
    return index in set(range(122, 127)) | set(range(177, 183)) | set(range(264, 270))


def build():
    data = json.loads(SOURCE.read_text())
    paras = data["paragraphs"]
    doc = Document()
    configure_styles(doc)
    bullet_num, decimal_num = add_numbering(doc)
    update_fields_on_open(doc)

    cover_sec = doc.sections[0]; section_geometry(cover_sec); set_header(cover_sec); set_footer(cover_sec, numbered=False)
    cover(doc)

    pre = doc.add_section(WD_SECTION_START.NEW_PAGE); section_geometry(pre); set_page_numbering(pre, "lowerRoman", 1); set_header(pre); set_footer(pre, numbered=True)
    bookmark_id = 100
    for i in range(10, 25):
        text = paras[i]["text"].strip()
        if not text: continue
        if paras[i].get("namedStyleType") == "HEADING_1":
            add_heading(doc, text, 1, front=True, anchor=f"front_{i}", bookmark_id=bookmark_id); bookmark_id += 1
        elif i == 11:
            for _ in range(5): doc.add_paragraph()
            add_body(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        elif i in (19, 24):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before = Pt(8)
            label, value = text.split(":", 1)
            r = p.add_run(label + " :"); set_font(r, 10, bold=True, color=ACCENT)
            r = p.add_run(value); set_font(r, 10, italic=True, color=INK)
        else:
            add_body(doc, text)

    body_pages = {
        118: 1, 121: 1, 127: 1, 129: 2, 130: 2, 132: 2, 134: 2, 139: 2, 155: 3,
        157: 4, 158: 4, 160: 4, 163: 4, 165: 4, 167: 5, 168: 5, 170: 5, 183: 5,
        197: 7, 205: 8, 207: 9, 209: 9, 211: 10, 212: 10, 214: 10, 216: 10,
        219: 10, 247: 14, 254: 15, 256: 15, 258: 16, 259: 16, 261: 16, 263: 16,
        270: 17, 280: 18, 281: 18, 287: 18, 289: 18,
    }
    toc_entries = [
        ("I. Dédicace", "i", "front_10", 1), ("II. Remerciements", "ii", "front_12", 1),
        ("III. Résumé", "iii", "front_15", 1), ("IV. Abstract", "iv", "front_20", 1),
        ("V. Table des matières", "v", "front_toc", 1), ("VI. Liste des figures", "vi", "front_figures", 1),
        ("VII. Liste des tableaux", "vii", "front_tables", 1), ("VIII. Liste des abréviations", "viii", "front_abbr", 1),
    ]
    for i, page in body_pages.items():
        level = 1 if paras[i].get("namedStyleType") == "HEADING_1" else 2
        toc_entries.append((remap_heading(paras[i]["text"].strip()), page, f"body_{i}", level))
    toc_entries.append(("Annexe D — Commandes de reproduction", 18, "annex_d", 2))

    add_heading(doc, "V. Table des matières", 1, front=True, anchor="front_toc", bookmark_id=bookmark_id); bookmark_id += 1
    add_live_index(doc, toc_entries, 'TOC \\o "1-3" \\h \\z \\u')
    add_heading(doc, "VI. Liste des figures", 1, front=True, anchor="front_figures", bookmark_id=bookmark_id); bookmark_id += 1
    figure_entries = [
        ("Figure 1.1 — Planification du projet par sprints", 3, "fig_1_1", 1),
        ("Figure 3.1 — Cas d’utilisation de DataWatch", 6, "fig_3_1", 1),
        ("Figure 3.2 — Séquence de détection et de traitement d’un incident", 7, "fig_3_2", 1),
        ("Figure 3.3 — Diagramme de classes métier simplifié", 7, "fig_3_3", 1),
        ("Figure 3.4 — Architecture en couches de DataWatch", 8, "fig_3_4", 1),
        ("Figure 4.1 — Vue Opérations et file d’incidents prioritaires", 11, "fig_4_1", 1),
        ("Figure 4.2 — Détail de l’incident critique orders", 12, "fig_4_2", 1),
        ("Figure 4.3 — Profil et métriques de la table orders", 12, "fig_4_3", 1),
        ("Figure 4.4 — Route d’alerte e-mail configurée pour la démonstration", 13, "fig_4_4", 1),
        ("Figure 4.5 — File de travail de gouvernance IA en mode observe-only", 13, "fig_4_5", 1),
        ("Figure 4.6 — Parcours de démonstration de l’incident orders", 14, "fig_4_6", 1),
        ("Figure 4.7 — Synthèse des validations locales et de leurs limites", 15, "fig_4_7", 1),
    ]
    add_live_index(doc, figure_entries, 'TOC \\h \\z \\t "Caption,1"')
    add_heading(doc, "VII. Liste des tableaux", 1, front=True, anchor="front_tables", bookmark_id=bookmark_id); bookmark_id += 1
    add_body(doc, "Aucun tableau n’est référencé dans le corps du rapport. Les comparaisons et résultats sont présentés sous forme de synthèses narratives ou de figures afin de préserver la lisibilité.", italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_heading(doc, "VIII. Liste des abréviations", 1, front=True, anchor="front_abbr", bookmark_id=bookmark_id); bookmark_id += 1
    add_abbreviations(doc)

    body = doc.add_section(WD_SECTION_START.NEW_PAGE); section_geometry(body); set_page_numbering(body, "decimal", 1); set_header(body, "DataWatch — Rapport de Projet de Fin d’Études"); set_footer(body, numbered=True)

    extras = {
        "3.3 Originalité de la solution": [
            "Le DSL de moniteurs complète cette chaîne par des définitions strictes, versionnées et liées au schéma observé. Une révision est validée puis prévisualisée avant activation ; l’exécution utilise un plan déterministe et un état d’évaluation séparé de l’historique immuable. Ce choix évite qu’une modification silencieuse change le comportement d’un contrôle déjà audité.",
        ],
        "4.4 Architecture proposée": [
            "Le cycle de profilage démarre dans APScheduler, qui place une tâche dans Redis. Celery déchiffre la configuration de la source dans le contexte de l’organisation, construit une requête agrégée, persiste le profil, puis déclenche les détecteurs et les moniteurs personnalisés. La narration et l’alerte sont exécutées après la persistance de l’incident afin qu’un échec externe ne supprime pas le fait technique observé.",
        ],
        "4.5 Modèle de données et sécurité": [
            "La sécurité repose également sur des clés étrangères composites qui prouvent l’appartenance des entités à une même organisation. Les révisions, résultats terminaux et preuves de gouvernance sont append-only lorsque leur valeur d’audit doit être conservée. Les clés API sont stockées sous forme de hachage et les secrets de connexion sont chiffrés avec une clé dérivée par organisation.",
        ],
        "5.3 Technologies utilisées": [
            "Le choix des technologies répond à une logique de responsabilités clairement séparées : FastAPI pour les contrats et l’asynchronisme, PostgreSQL pour les contraintes transactionnelles et l’historique, Redis et Celery pour les travaux différés, React pour l’investigation visible, et Docker Compose pour reproduire l’environnement de démonstration. Cette organisation facilite l’évolution indépendante des composants tout en conservant un parcours de bout en bout testable.",
        ],
    }

    for i in range(118, len(paras)):
        if i in SKIP_BODY_INDEXES: continue
        text = paras[i]["text"].strip()
        if not text: continue
        style = paras[i].get("namedStyleType")
        if i == 146:
            text = "Le planning associe chaque incrément à un objectif vérifiable et à un livrable exploitable dans la démonstration. La figure 1.1 synthétise les sept sprints, leurs chevauchements et leurs principaux jalons."
        if i in range(147, 152):
            continue
        fig = re.match(r"^(Figure \d+\.\d+)\s*[-–—]\s*(.+?)\.?$", text)
        if fig:
            label, caption = fig.group(1), fig.group(2)
            source_label = label
            if label == "Figure 4.7": label = "Figure 4.6"
            elif label == "Figure 4.6": label = "Figure 4.7"
            add_picture(doc, FIGURES[source_label], label, caption, anchor="fig_" + label.split()[1].replace(".", "_"), bookmark_id=bookmark_id); bookmark_id += 1
            continue
        if style in ("HEADING_1", "HEADING_2", "HEADING_3"):
            level = {"HEADING_1": 1, "HEADING_2": 2, "HEADING_3": 3}[style]
            anchor = f"body_{i}" if i in body_pages else None
            add_heading(doc, remap_heading(text), level, anchor=anchor, bookmark_id=bookmark_id if anchor else None)
            if anchor: bookmark_id += 1
            for extra in extras.get(text, []): add_body(doc, extra)
            continue
        if i in range(271, 280):
            p = add_body(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, after=5)
            p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)
            for run in p.runs: set_font(run, 9.7, color=INK)
            continue
        if is_bullet_index(i):
            add_list(doc, text, bullet_num)
            continue
        if i in range(282, 287):
            add_list(doc, text, decimal_num)
            continue
        if i == 249:
            p = add_body(doc, text.upper(), align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=6)
            p.paragraph_format.space_before = Pt(6)
            paragraph_border(p, color=ACCENT, size="6", space="4")
            continue
        add_body(doc, text)

    add_heading(doc, "Annexe D — Commandes de reproduction", 2, anchor="annex_d", bookmark_id=bookmark_id); bookmark_id += 1
    add_body(doc, "Les commandes suivantes reconstruisent la pile de démonstration, réinitialisent les données et régénèrent les cinq captures utilisées dans le rapport.")
    add_code_block(doc, [
        "docker compose up -d --wait",
        "docker compose --profile seed run --rm --entrypoint python seed /scripts/quickstart.py --reset",
        "curl -fsS http://localhost:8000/ready",
        "cd frontend && npm run capture:pfe",
    ])

    core = doc.core_properties
    core.title = "Rapport PFE — DataWatch"
    core.subject = "Conception et réalisation d’une plateforme SaaS de surveillance de la qualité des données"
    core.author = "Mounir Gaiby"
    core.keywords = "DataWatch, qualité des données, SaaS, intelligence artificielle, ISGA"
    core.comments = "Rapport de Projet de Fin d’Études — ISGA Casablanca"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
